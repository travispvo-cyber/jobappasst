"""Extract text from PDF and DOCX resume files"""

import pdfplumber
from pathlib import Path
from docx import Document
from typing import Optional


def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extract text from a PDF file.

    Args:
        file_path: Path to the PDF file

    Returns:
        str: Extracted text from the PDF

    Raises:
        FileNotFoundError: If the file doesn't exist
        Exception: If PDF parsing fails
    """
    if not file_path.exists():
        raise FileNotFoundError(f"PDF file not found: {file_path}")

    try:
        text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)

        extracted = "\n\n".join(text)

        if not extracted.strip():
            raise Exception("No text could be extracted from the PDF")

        return extracted

    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        str: Extracted text from the DOCX

    Raises:
        FileNotFoundError: If the file doesn't exist
        Exception: If DOCX parsing fails
    """
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    try:
        doc = Document(file_path)
        text = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)

        extracted = "\n".join(text)

        if not extracted.strip():
            raise Exception("No text could be extracted from the DOCX")

        return extracted

    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")


def extract_text_from_resume(file_path: str | Path) -> str:
    """
    Extract text from a resume file (PDF or DOCX).

    Automatically detects file type based on extension.

    Args:
        file_path: Path to the resume file (PDF or DOCX)

    Returns:
        str: Extracted text from the resume

    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file doesn't exist
        Exception: If text extraction fails
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    # Get file extension (lowercase)
    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif extension in [".docx", ".doc"]:
        # Note: .doc (old Word format) requires docx, which may not work perfectly
        # For best results, convert .doc to .docx
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type: {extension}. "
            "Supported types: .pdf, .docx"
        )


# Example usage and testing
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python resume_parser.py <path_to_resume>")
        sys.exit(1)

    resume_path = sys.argv[1]

    try:
        text = extract_text_from_resume(resume_path)
        print("=" * 60)
        print("EXTRACTED TEXT:")
        print("=" * 60)
        print(text)
        print("=" * 60)
        print(f"\nTotal characters: {len(text)}")
        print(f"Total lines: {len(text.splitlines())}")

    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)
